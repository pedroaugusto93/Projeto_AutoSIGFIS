# make_readme_docx.py
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("Falta a biblioteca 'python-docx'. Instale com: pip install python-docx")
    sys.exit(1)

OUTPUT_PATH = Path("README_AutoSIGFIS2.docx")  # ajuste se quiser outro local/nome

def add_title(doc, text):
    p = doc.add_heading(text, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_h1(doc, text):
    doc.add_heading(text, level=1)

def add_h2(doc, text):
    doc.add_heading(text, level=2)

def add_paragraph(doc, text):
    doc.add_paragraph(text)

def add_code_block(doc, text):
    """Simula bloco de código com fonte monoespaçada."""
    for line in text.strip("\n").split("\n"):
        run = doc.add_paragraph().add_run(line)
        run.font.name = "Consolas"
        # Garantir fonte em alguns ambientes
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        run.font.size = Pt(10)

def main():
    doc = Document()

    # Título
    add_title(doc, "Projeto AutoSIGFIS / AutoSIGFIS Project")

    # ---------------------------- Português ----------------------------
    add_h1(doc, "Descrição (Português)")
    add_paragraph(
        doc,
        "Automação em Python + Selenium para preenchimento de cadastros no sistema SIGFIS "
        "(Sistema de Gestão de Fiscalização). O sistema lê dados de uma planilha Excel "
        "(cadastro.xlsx) e realiza o preenchimento automático das abas: Dados Básicos, Itens, "
        "Documentos e Empenhos."
    )

    add_h2(doc, "Estrutura do Projeto")
    add_code_block(doc, r"""
Projeto_AutoSIGFIS/
├── config.py             # Configurações fixas e leitura da planilha
├── helpers.py            # Funções utilitárias (inputs, selects, waits…)
├── main.py               # Fluxo principal da automação
├── page_dados_basicos.py # Lógica da aba "Dados Básicos"
├── page_itens.py         # Lógica da aba "Itens"
├── page_empenhos.py      # Lógica da aba "Empenhos"
├── page_documentos.py    # Lógica da aba "Documentos"
├── cadastro.xlsx         # Planilha de entrada
└── requirements.txt      # Dependências do projeto
""")

    add_h2(doc, "Instalação")
    add_code_block(doc, r"""
# 1) Clone o repositório
git clone https://github.com/seu-usuario/Projeto_AutoSIGFIS.git
cd Projeto_AutoSIGFIS

# 2) Crie e ative a venv
python -m venv venv
venv\Scripts\Activate.ps1   # Windows PowerShell
# ou
source venv/bin/activate    # Linux/Mac

# 3) Instale as dependências
pip install -r requirements.txt
""")

    add_h2(doc, "Execução")
    add_paragraph(doc, "1) Inicie o Google Chrome em modo depuração:")
    add_code_block(doc, r"""
"C:\Program Files\Google\Chrome\Application\chrome.exe" --remote-debugging-port=9222 --user-data-dir="C:\ChromeDebugProfile"
""")
    add_paragraph(doc, "2) Execute o script:")
    add_code_block(doc, "python main.py")

    add_h2(doc, "Tecnologias")
    add_paragraph(doc, "Python 3.9+, Selenium, Pandas, OpenPyXL, WebDriver Manager")

    add_h2(doc, "Autor")
    add_paragraph(doc, "Pedro - Projeto desenvolvido para automatizar cadastros no SIGFIS e compor portfólio público.")

    # ---------------------------- Inglês ----------------------------
    add_h1(doc, "Description (English)")
    add_paragraph(
        doc,
        "Automation in Python + Selenium for filling out records in the SIGFIS system. "
        "The bot reads data from an Excel file (cadastro.xlsx) and automatically fills in the tabs: "
        "Basic Data, Items, Documents and Commitments."
    )

    add_h2(doc, "Project Structure")
    add_code_block(doc, r"""
Projeto_AutoSIGFIS/
├── config.py             # Fixed settings and Excel loading
├── helpers.py            # Utility functions (inputs, selects, waits…)
├── main.py               # Main automation flow
├── page_dados_basicos.py # "Basic Data" tab logic
├── page_itens.py         # "Items" tab logic
├── page_empenhos.py      # "Commitments" tab logic
├── page_documentos.py    # "Documents" tab logic
├── cadastro.xlsx         # Input spreadsheet
└── requirements.txt      # Project dependencies
""")

    add_h2(doc, "Installation")
    add_code_block(doc, r"""
# 1) Clone repository
git clone https://github.com/your-user/Projeto_AutoSIGFIS.git
cd Projeto_AutoSIGFIS

# 2) Create and activate virtual environment
python -m venv venv
venv\Scripts\Activate.ps1   # Windows PowerShell
# or
source venv/bin/activate    # Linux/Mac

# 3) Install dependencies
pip install -r requirements.txt
""")

    add_h2(doc, "Execution")
    add_paragraph(doc, "1) Launch Google Chrome in debug mode:")
    add_code_block(doc, r"""
"C:\Program Files\Google\Chrome\Application\chrome.exe" --remote-debugging-port=9222 --user-data-dir="C:\ChromeDebugProfile"
""")
    add_paragraph(doc, "2) Run the script:")
    add_code_block(doc, "python main.py")

    add_h2(doc, "Technologies")
    add_paragraph(doc, "Python 3.9+, Selenium, Pandas, OpenPyXL, WebDriver Manager")

    add_h2(doc, "Author")
    add_paragraph(doc, "Pedro - Project developed to automate SIGFIS records and serve as a public portfolio.")

    # Salvar
    try:
        OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
        doc.save(OUTPUT_PATH)
        print(f"OK! Arquivo gerado em:\n{OUTPUT_PATH.resolve()}")
    except PermissionError:
        print(f"Feche o arquivo se ele estiver aberto e rode novamente:\n{OUTPUT_PATH.resolve()}")
    except Exception as e:
        print("Erro ao salvar o DOCX:", e)
        raise

if __name__ == "__main__":
    main()
# make_readme_docx.py
# Gerador de README.md para DOCX    